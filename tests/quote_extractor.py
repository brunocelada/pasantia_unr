import os
import logging
from docx import Document
from docx.shared import Pt
from datetime import datetime
import sys

# Configurar logging (archivo y consola, consola sobrescribe)
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def log_sobrescribir(mensaje):
    sys.stdout.write("\r" + " " * 120 + "\r")  # Borrar línea previa
    sys.stdout.write(mensaje)
    sys.stdout.flush()

def solicitar_direccion_carpeta(mensaje):
    ruta = input(mensaje)
    if not os.path.isdir(ruta):
        raise NotADirectoryError(f"La ruta proporcionada no es válida: {ruta}")
    return os.path.abspath(ruta)

def buscar_archivos_log(carpeta_base):
    archivos_log = []
    for root, _, files in os.walk(carpeta_base):
        for file in files:
            if file.endswith(".log"):
                archivos_log.append(os.path.join(root, file))
    return archivos_log

def verificar_terminacion_correcta(path):
    try:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()
        for line in reversed(lines[-50:]):
            if "Normal termination" in line:
                return True, lines
        return False, lines
    except Exception as e:
        logging.error(f"Error leyendo archivo {path}: {e}")
        return False, []

def extraer_cita(lines):
    try:
        end_index = None
        for i in range(len(lines)-1, -1, -1):
            if "Normal termination" in lines[i]:
                end_index = i
                break
        if end_index is None:
            return None, None

        # Buscar hacia arriba desde la terminación hasta encontrar '@'
        start_index = None
        for i in range(end_index, -1, -1):
            if "@" in lines[i]:
                start_index = i
                break
        if start_index is None:
            return None, None

        # Extraer líneas entre @ y terminación
        quote_block = lines[start_index+1:end_index]

        # Limpiar líneas irrelevantes
        cleaned_lines = []
        for line in quote_block:
            line = line.strip()
            if not line:
                continue
            if line.startswith("The archive entry for this job was punched."):
                continue
            if line.startswith("Job cpu time") or line.startswith("Elapsed time") or line.startswith("File lengths"):
                continue
            cleaned_lines.append(line)

        return ('\n'.join(cleaned_lines).strip() if cleaned_lines else None), None
    except Exception as e:
        logging.error(f"Error extrayendo cita: {e}")
        return None, None

def guardar_citas_en_docx(citas, base_name, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    doc_count = 1
    for i in range(0, len(citas), 100):
        doc = Document()
        chunk = citas[i:i+100]
        for cita, fuente in chunk:
            p = doc.add_paragraph()
            run = p.add_run(cita + "\n")
            run.font.size = Pt(11)
            if fuente:
                p_fuente = doc.add_paragraph()
                fuente_run = p_fuente.add_run(f"({fuente})")
                fuente_run.font.size = Pt(9)
        filename = os.path.join(output_dir, f"{base_name}_{doc_count:02}.docx")
        doc.save(filename)
        logging.info(f"Guardado: {filename}")
        doc_count += 1

def main():
    try:
        carpeta_busqueda = solicitar_direccion_carpeta("Ingrese la carpeta a recorrer: ")
        carpeta_salida = solicitar_direccion_carpeta("Ingrese la carpeta donde guardar los DOCX: ")
        nombre_base = input("Ingrese el nombre base para los archivos Word: ").strip()

        archivos_log = buscar_archivos_log(carpeta_busqueda)
        logging.info(f"Se encontraron {len(archivos_log)} archivos .log")

        citas = []
        for idx, archivo in enumerate(archivos_log, 1):
            log_sobrescribir(f"Procesando archivo {idx}/{len(archivos_log)}: {os.path.basename(archivo)}")
            terminado, lines = verificar_terminacion_correcta(archivo)
            if not terminado:
                continue
            cita, _ = extraer_cita(lines)
            if cita:
                citas.append((cita, os.path.basename(archivo)))

        print()  # Para dejar espacio después de la última sobrescritura
        if citas:
            guardar_citas_en_docx(citas, nombre_base, carpeta_salida)
        else:
            logging.info("No se encontraron citas válidas para guardar.")

    except Exception as e:
        logging.error(f"Error general: {e}")

if __name__ == "__main__":
    main()
